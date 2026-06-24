"""
Genera un file .docx con la stima costi e tempi per:
1. Sviluppo dello stato attuale dell'applicativo (junior developer)
2. Sviluppo delle 4 fasi della roadmap futura (Audit log, PDF, Altri Cantoni, OCR)

Output: G:\\DdC\\docs\\stima-junior-developer.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(r"G:\DdC\docs\stima-junior-developer.docx")

# ── Tariffe di riferimento Svizzera, junior 0-2 anni ──────────────────────────
RATE_INTERNAL_LOW = 55   # CHF/h costo aziendale interno (stipendio + oneri)
RATE_INTERNAL_HIGH = 65
RATE_EXTERNAL_LOW = 90   # CHF/h consulenza esterna
RATE_EXTERNAL_HIGH = 110


def chf(amount: float) -> str:
    return f"CHF {amount:,.0f}".replace(",", "'")


def cost_range(hours_low: int, hours_high: int) -> str:
    low = hours_low * RATE_INTERNAL_LOW
    high = hours_high * RATE_EXTERNAL_HIGH
    return f"{chf(low)} – {chf(high)}"


# ── Stima stato attuale ───────────────────────────────────────────────────────
CURRENT_STATE = [
    ("Studio normativa fiscale (Art. 25 LT, Art. 26 LIFD, RS 642.118.1, accordo frontalieri CH-IT)", 40),
    ("Setup progetto FastAPI + Pydantic + uv + pyproject", 8),
    ("Schemi request/response Pydantic con validazioni", 16),
    ("Loader YAML + modelli regole fiscali (rules/models.py)", 20),
    ("Motore cantonale TI — trasporto (auto, mezzi pubblici, bici, misto, moto)", 50),
    ("Motore federale IFD + cap CHF 3'300 + RS 642.118.1", 25),
    ("Motore pasti (Art. 6 + Art. 9, tutte le situazioni)", 25),
    ("Modulo altre spese 3% / forfait IC", 16),
    ("Casi speciali — frontalieri, residenti settimanali, home office, turni", 25),
    ("Lohnausweis — campi D, F, cifra 11, cifra 13.2.2 + logica di precedenza", 25),
    ("Geocoding — swisstopo (primario) + Nominatim (fallback) + calcolo distanza", 35),
    ("Endpoint API v1 (/calculate, /rules/{year}, /health) + main.py + lifespan", 16),
    ("Security hardening — rate limit, API key, security headers, body size, CORS", 35),
    ("Frontend statico HTML + CSS + JS (form, validazione, render risultati)", 35),
    ("Test unit (8 file, ~90 test su engine e validazioni)", 50),
    ("Test integration (API, security, headers)", 25),
    ("Docker + deploy VPS Infomaniak Nano + systemd", 25),
    ("Documentazione (README, deploy-infomaniak-nano.md, security.md)", 15),
    ("Debug, iterazioni, fix bug (buffer 25%)", 110),
]
CURRENT_TOTAL_HOURS = sum(h for _, h in CURRENT_STATE)

# ── Stima 4 fasi roadmap futura ───────────────────────────────────────────────
PHASE_1 = [
    ("Setup modulo app/audit/ + interfaccia AuditStore (Protocol)", 4),
    ("Schema SQLite + serializzazione JSON request/response", 8),
    ("Hook nell'endpoint /calculate + nuovo endpoint GET /audit/{request_id}", 6),
    ("Tests integration (save + retrieve + audit_enabled=False)", 6),
    ("Buffer iterazioni 25%", 6),
]
PHASE_2 = [
    ("Setup modulo app/reports/ + dipendenza WeasyPrint", 4),
    ("Template HTML + CSS print-friendly (@page A4, margini, tipografia)", 16),
    ("Renderer PDF + endpoint GET /v1/deduction/{request_id}/pdf", 8),
    ("Test integration (content-type PDF, header %PDF-)", 6),
    ("Iterazioni layout su carta A4 reale", 8),
    ("Buffer 25%", 10),
]
PHASE_3 = [
    ("Studio normativa StGB Zurigo + StG Berna (Steueramt cantonali)", 30),
    ("Refactor rules/models.py (cantonal dict-keyed) + loader.py", 12),
    ("Refactor cantonal_engine.py + calculator.py (rimozione cantonal_TI hardcoded)", 16),
    ("Migrazione legal_reference da Python a YAML", 8),
    ("Aggiunta CantonCode enum + propagazione in DeductionRequest", 6),
    ("Scrittura rules/2026.yaml per ZH + BE (trasporto + pasti)", 12),
    ("Tests parametrizzati ZH/BE + nuovi casi", 20),
    ("Verifica regressione TI (zero modifica risultati esistenti)", 6),
    ("Buffer 30% (refactor invasivo)", 30),
]
PHASE_4 = [
    ("Studio layout standardizzato BFS Lohnausweis (coordinate campi)", 12),
    ("Setup pdfplumber + struttura modulo app/ocr/", 8),
    ("Parser PDF testuale (estrazione per coordinate dei campi D/F/11/13.2.2)", 25),
    ("Fallback OCR Tesseract per scansioni immagine (pytesseract)", 12),
    ("Confidence scoring per ogni campo estratto", 6),
    ("Endpoint POST /v1/lohnausweis/parse + UploadFile + validazione", 8),
    ("Schema Pydantic LohnausweisFields + integrazione opzionale con /calculate", 6),
    ("Test con fixture PDF reali anonimizzati (3 campioni: moderno, vecchio, scansione)", 16),
    ("Buffer 30% (alta variabilità OCR su scansioni reali)", 27),
]

PHASES = [
    ("Fase 1 — Audit log persistente (SQLite)", PHASE_1, "Bassa"),
    ("Fase 2 — Export PDF del calcolo (WeasyPrint)", PHASE_2, "Media"),
    ("Fase 3 — Estensione altri Cantoni (ZH + BE)", PHASE_3, "Alta"),
    ("Fase 4 — OCR Lohnausweis upload", PHASE_4, "Alta"),
]


# ── Document generation ───────────────────────────────────────────────────────

def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_para(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_breakdown_table(doc: Document, rows: list[tuple[str, int]]) -> None:
    table = doc.add_table(rows=1 + len(rows) + 1, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0].cells
    header[0].text = "Attività"
    header[1].text = "Ore stimate"
    for cell in header:
        cell.paragraphs[0].runs[0].bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, (label, hours) in enumerate(rows, start=1):
        row = table.rows[i].cells
        row[0].text = label
        row[1].text = str(hours)
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    total = sum(h for _, h in rows)
    last = table.rows[-1].cells
    last[0].text = "Totale"
    last[1].text = str(total)
    for cell in last:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    last[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Larghezza colonne
    for row in table.rows:
        row.cells[0].width = Cm(13)
        row.cells[1].width = Cm(3)


def add_cost_table(doc: Document, hours_low: int, hours_high: int) -> None:
    table = doc.add_table(rows=5, cols=4)
    table.style = "Light Grid Accent 1"

    headers = ["Profilo", "Tariffa CHF/h", "Costo a ore basse", "Costo a ore alte"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    rows_data = [
        ("Junior interno (lower)", RATE_INTERNAL_LOW, hours_low * RATE_INTERNAL_LOW, hours_high * RATE_INTERNAL_LOW),
        ("Junior interno (upper)", RATE_INTERNAL_HIGH, hours_low * RATE_INTERNAL_HIGH, hours_high * RATE_INTERNAL_HIGH),
        ("Junior esterno (lower)", RATE_EXTERNAL_LOW, hours_low * RATE_EXTERNAL_LOW, hours_high * RATE_EXTERNAL_LOW),
        ("Junior esterno (upper)", RATE_EXTERNAL_HIGH, hours_low * RATE_EXTERNAL_HIGH, hours_high * RATE_EXTERNAL_HIGH),
    ]
    for i, (profile, rate, low_cost, high_cost) in enumerate(rows_data, start=1):
        row = table.rows[i].cells
        row[0].text = profile
        row[1].text = chf(rate)
        row[2].text = chf(low_cost)
        row[3].text = chf(high_cost)
        for j in (1, 2, 3):
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title page
title = doc.add_heading("Stima costi e tempistiche", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = subtitle.add_run(
    "Web Service Deduzioni Trasferta Casa-Lavoro — Canton Ticino + IFD\n"
    "Profilo di riferimento: Junior Developer (0-2 anni esperienza)"
)
sub.italic = True
sub.font.size = Pt(12)

doc.add_paragraph()

# ── Premessa metodologica ─────────────────────────────────────────────────────
add_heading(doc, "1. Premessa metodologica", level=1)
add_para(
    doc,
    "Le stime sono espresse in ore di sviluppo nette e includono un buffer per debug, "
    "iterazioni e code review. Il profilo di riferimento è un developer junior (0-2 anni "
    "di esperienza professionale) che lavora autonomamente, senza supporto continuativo di "
    "un senior, su tecnologie comuni (FastAPI, Pydantic, SQL, HTML/CSS/JS). Le ore di "
    "studio della normativa fiscale svizzera sono incluse, perché un junior tipico non "
    "ha conoscenza pregressa di LT, LIFD e ordinanze DFF.",
)

add_para(doc, "Tariffe di riferimento (mercato svizzero, 2026):", bold=True)
add_para(doc, f"  • Junior interno (stipendio + oneri sociali): CHF {RATE_INTERNAL_LOW}–{RATE_INTERNAL_HIGH}/h")
add_para(doc, f"  • Junior in consulenza esterna: CHF {RATE_EXTERNAL_LOW}–{RATE_EXTERNAL_HIGH}/h")
add_para(
    doc,
    "Una giornata lavorativa è considerata 8 ore effettive. Una settimana = 5 giornate = 40 ore. "
    "Le stime non includono costi di infrastruttura (VPS Infomaniak ~CHF 5-6/mese), "
    "licenze software, formazione, o gestione del progetto (PM, QA dedicato).",
)

# ── Sezione 2: Stato attuale ──────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "2. Stato attuale dell'applicativo", level=1)
add_para(
    doc,
    "Il servizio attualmente in produzione su VPS Infomaniak Nano implementa il calcolo "
    "completo delle deduzioni trasferta per il Canton Ticino (Art. 25 LT) e per l'imposta "
    "federale diretta (Art. 26 LIFD + RS 642.118.1), con supporto multi-anno (regole 2025 "
    "e 2026), tutti i mezzi di trasporto, gestione casi speciali (frontalieri, residenti "
    "settimanali, home office, turni) e integrazione completa dei campi Lohnausweis "
    "(D, F, cifra 11, cifra 13.2.2). Hardening di sicurezza completato (93/93 test), "
    "frontend statico funzionante, deploy Docker + systemd documentato.",
)

add_heading(doc, "2.1 Breakdown delle attività", level=2)
add_breakdown_table(doc, CURRENT_STATE)

doc.add_paragraph()
add_heading(doc, "2.2 Tempistiche stato attuale", level=2)
total_low = int(CURRENT_TOTAL_HOURS * 0.85)
total_high = int(CURRENT_TOTAL_HOURS * 1.10)
add_para(doc, f"Totale ore stimate: {CURRENT_TOTAL_HOURS} h", bold=True)
add_para(doc, f"Range realistico (±15%): {total_low} – {total_high} h")
add_para(doc, f"Equivalenti in giornate (8h): {total_low // 8} – {total_high // 8} giornate")
add_para(doc, f"Equivalenti in settimane (40h): {total_low // 40} – {total_high // 40} settimane")
add_para(doc, f"Calendario realistico full-time: {total_low // 160} – {total_high // 160} mesi")

add_heading(doc, "2.3 Costi stato attuale", level=2)
add_cost_table(doc, total_low, total_high)
doc.add_paragraph()
add_para(
    doc,
    f"Range complessivo: {cost_range(total_low, total_high)} a seconda del profilo "
    "(interno vs esterno) e dell'efficienza effettiva sul progetto.",
    bold=True,
)

# ── Sezione 3: Nuove aggiunte ─────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "3. Roadmap nuove aggiunte", level=1)
add_para(
    doc,
    "Le 4 fasi previste dalla roadmap (file di piano: "
    "C:\\Users\\oudab\\.claude\\plans\\quali-prossime-funyionalit-a-potrei-wiggly-matsumoto.md) "
    "estendono il servizio con: audit log persistente, export PDF dei calcoli, supporto "
    "ad altri Cantoni svizzeri e OCR del Lohnausweis. Le fasi sono indipendenti e "
    "rilasciabili separatamente, con sequenza consigliata Audit → PDF → Cantoni → OCR.",
)

phase_totals: list[tuple[str, int, str]] = []
for phase_name, breakdown, risk in PHASES:
    add_heading(doc, phase_name, level=2)
    add_para(doc, f"Rischio implementativo: {risk}", bold=True)
    add_breakdown_table(doc, breakdown)
    total = sum(h for _, h in breakdown)
    phase_totals.append((phase_name, total, risk))
    doc.add_paragraph()

# ── Sezione 4: Riepilogo aggiunte ────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "4. Riepilogo nuove aggiunte", level=1)

table = doc.add_table(rows=1 + len(phase_totals) + 1, cols=4)
table.style = "Light Grid Accent 1"
headers = ["Fase", "Rischio", "Ore stimate", "Costo (range completo)"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

new_total_hours = 0
for i, (name, hours, risk) in enumerate(phase_totals, start=1):
    row = table.rows[i].cells
    row[0].text = name
    row[1].text = risk
    row[2].text = str(hours)
    low = hours * RATE_INTERNAL_LOW
    high = hours * RATE_EXTERNAL_HIGH
    row[3].text = f"{chf(low)} – {chf(high)}"
    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    new_total_hours += hours

last = table.rows[-1].cells
last[0].text = "Totale roadmap"
last[1].text = "—"
last[2].text = str(new_total_hours)
low_tot = new_total_hours * RATE_INTERNAL_LOW
high_tot = new_total_hours * RATE_EXTERNAL_HIGH
last[3].text = f"{chf(low_tot)} – {chf(high_tot)}"
for cell in last:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
last[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
last[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_paragraph()
new_low = int(new_total_hours * 0.85)
new_high = int(new_total_hours * 1.15)
add_para(doc, f"Totale ore roadmap (somma fasi): {new_total_hours} h", bold=True)
add_para(doc, f"Range realistico (±15%): {new_low} – {new_high} h")
add_para(doc, f"Equivalenti in giornate (8h): {new_low // 8} – {new_high // 8} giornate")
add_para(doc, f"Calendario full-time: {new_low / 160:.1f} – {new_high / 160:.1f} mesi")

# ── Sezione 5: Totale generale ───────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "5. Totale generale (stato attuale + roadmap)", level=1)
grand_low = total_low + new_low
grand_high = total_high + new_high
add_para(doc, f"Ore complessive: {grand_low} – {grand_high} h", bold=True, size=12)
add_para(
    doc,
    f"Costo complessivo, range completo da junior interno lower a junior esterno upper: "
    f"{cost_range(grand_low, grand_high)}",
    bold=True,
    size=12,
)

doc.add_paragraph()
add_heading(doc, "Note finali", level=2)
add_para(
    doc,
    "• Le stime presuppongono che la normativa fiscale per i nuovi Cantoni (Fase 3) e "
    "i campioni reali di Lohnausweis (Fase 4) vengano forniti dal committente. In caso "
    "contrario, aggiungere 20-40 h di ricerca per ciascuna fase.",
)
add_para(
    doc,
    "• Se il progetto fosse svolto da un mid-level developer (3-5 anni esperienza), il "
    "tempo totale si ridurrebbe del 30-40%, ma la tariffa oraria salirebbe del 40-60%. "
    "Il costo finale tende a essere simile, con minore rischio e tempi più certi.",
)
add_para(
    doc,
    "• Una manutenzione evolutiva post-rilascio realistica è ~10-15% delle ore di "
    "sviluppo iniziale per anno (es. aggiornamento regole fiscali ogni anno, fix sicurezza, "
    "compatibilità nuove versioni Python/FastAPI).",
)

# ── Save ──────────────────────────────────────────────────────────────────────
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(f"Generato: {OUTPUT}")
print(f"Ore stato attuale: {CURRENT_TOTAL_HOURS}")
print(f"Ore roadmap: {new_total_hours}")
print(f"Costo complessivo: {cost_range(total_low + new_low, total_high + new_high)}")
